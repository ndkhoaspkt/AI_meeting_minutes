from openai import OpenAI # OpenAI API key
from docx import Document # Word document
import speech_recognition as sr # Speech recognition
import subprocess # Subprocess

client = OpenAI( # OpenAI API key
    api_key="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",
)

def transcribe_audio(audio_file_path): # Transcribe the audio file 
    r = sr.Recognizer()
    try:
        with sr.AudioFile(audio_file_path) as source:
            audio_data = r.record(source)
            text = r.recognize_google(audio_data)
        return text
    except ValueError as e:
        print(f"Error transcribing audio file: {e}")
        return ""

def convert_to_pcm_wav(input_file, output_file): # Convert the audio file to PCM WAV format
    command = ['ffmpeg', '-i', input_file, '-acodec', 'pcm_s16le', '-ac', '1', '-ar', '16000', output_file]
    subprocess.run(command, check=True)

def abstract_summary_extraction(transcription): # Extract the abstract summary of the meeting minutes
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0.5,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def key_points_extraction(transcription): # Extract the key points of the meeting minutes
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0.5,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def action_item_extraction(transcription): # Extract the action items of the meeting minutes
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0.5,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def sentiment_analysis(transcription): # Perform sentiment analysis on the meeting minutes
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        temperature=0.5,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def meeting_minutes(transcription): # Generate the meeting minutes
    abstract_summary = abstract_summary_extraction(transcription) # Summary of the meeting minutes
    key_points = key_points_extraction(transcription) # Key points of the meeting minutes
    action_items = action_item_extraction(transcription) # Action items of the meeting minutes
    sentiment = sentiment_analysis(transcription) #Sentiment analysis of the meeting minutes
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def save_as_docx(minutes, filename): # Save the meeting minutes as a Word document
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

audio_file_path = "Earningscall.wav" # Audio file path
pcm_wav_file_path = "Earningscall_pcm.wav" # PCM WAV file path
convert_to_pcm_wav(audio_file_path, pcm_wav_file_path) # Convert the audio file to PCM WAV format
transcription = transcribe_audio(pcm_wav_file_path) # Transcribe the audio file

if transcription: # If the transcription succeeded, generate the meeting minutes   
    minutes = meeting_minutes(transcription)
    print(minutes)
    save_as_docx(minutes, 'meeting_minutes.docx')
else: # If the transcription failed, print an error message
    print("Transcription failed. No meeting minutes were generated.")